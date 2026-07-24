"""
Экспорт отчёта сравнения LLM-моделей в DOCX.

Раньше был PPTX — презентация принудительно режется на слайды, из-за чего
таблицы и абзацы обрывались посередине без веской причины (плохая вёрстка).
Word не ограничен слайдами: рендерим Markdown отчёта судьи как обычный
структурированный документ — настоящие заголовки, таблицы, списки, без
искусственного разбиения на страницы.
"""
from __future__ import annotations

import io
import re
from datetime import datetime
from typing import Any

from agents.model_bench import compute_stats

_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")


def _strip_inline_code(text: str) -> str:
    return re.sub(r"`([^`]+)`", r"\1", text)


def _add_runs(paragraph, text: str) -> None:
    """**bold** — жирными runs, остальное обычным текстом (один paragraph)."""
    pos = 0
    for m in _BOLD_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        paragraph.add_run(m.group(1)).bold = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _split_table_row(line: str) -> list[str]:
    trimmed = line.strip().strip("|")
    return [c.strip() for c in trimmed.split("|")]


def _is_table_separator(cells: list[str]) -> bool:
    return len(cells) > 1 and all(re.fullmatch(r":?-{3,}:?", c) for c in cells if c)


def _add_markdown(doc, text: str) -> None:
    """Дописывает Markdown-текст (отчёт судьи) в документ как настоящие
    заголовки/таблицы/списки/абзацы — построчный разбор того же класса,
    что и NotionRenderer на фронте, только на выходе docx-элементы вместо React."""
    from docx.shared import Pt

    lines = text.strip("\n").splitlines()
    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()

        if not line:
            i += 1
            continue

        # Таблица: строка с | и следующая строка-разделитель ---|---
        if i + 1 < len(lines) and "|" in line:
            cells = _split_table_row(line)
            sep_cells = _split_table_row(lines[i + 1])
            if len(cells) > 1 and _is_table_separator(sep_cells):
                headers = [_strip_inline_code(c) for c in cells]
                rows: list[list[str]] = []
                j = i + 2
                while j < len(lines) and "|" in lines[j].strip():
                    rows.append([_strip_inline_code(c) for c in _split_table_row(lines[j])])
                    j += 1
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = "Light Grid Accent 1"
                for idx, h in enumerate(headers):
                    table.rows[0].cells[idx].paragraphs[0].add_run(h).bold = True
                for row_cells in rows:
                    out_cells = table.add_row().cells
                    for idx in range(len(headers)):
                        out_cells[idx].text = row_cells[idx] if idx < len(row_cells) else ""
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            for run in p.runs:
                                run.font.size = Pt(10)
                doc.add_paragraph()
                i = j
                continue

        header_m = re.match(r"^(#{1,6})\s+(.+)$", line)
        if header_m:
            level = min(len(header_m.group(1)), 4)
            doc.add_heading(_strip_inline_code(header_m.group(2)), level=level)
            i += 1
            continue

        if re.match(r"^[-*]\s+", line):
            p = doc.add_paragraph(style="List Bullet")
            _add_runs(p, _strip_inline_code(re.sub(r"^[-*]\s+", "", line)))
            i += 1
            continue

        if re.match(r"^\d+\.\s+", line):
            p = doc.add_paragraph(style="List Number")
            _add_runs(p, _strip_inline_code(re.sub(r"^\d+\.\s+", "", line)))
            i += 1
            continue

        if re.fullmatch(r"-{3,}", line) or re.fullmatch(r"={3,}", line):
            i += 1
            continue

        p = doc.add_paragraph()
        _add_runs(p, _strip_inline_code(line))
        i += 1


def build_docx(session: dict[str, Any]) -> bytes:
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        raise ValueError("DOCX не сформирован: на сервере не установлен python-docx")

    stats = compute_stats(session.get("targets", []))

    doc = Document()

    doc.add_heading("Сравнение LLM-моделей", level=0)
    best_p, best_m = session.get("best_provider", ""), session.get("best_model", "")
    created = str(session.get("created_at", ""))[:10]
    meta_lines = [f"Сформировано: {datetime.now().strftime('%d.%m.%Y')}"]
    if created:
        meta_lines.append(f"Сессия от {created}")
    if best_p:
        meta_lines.append(f"Лучшая модель: {best_p}/{best_m}")
    doc.add_paragraph(" · ".join(meta_lines)).runs[0].italic = True

    doc.add_heading("Условия замера", level=1)
    doc.add_paragraph(f"Промпт (первые 300 симв.): {session.get('prompt', '')[:300]}")
    doc.add_paragraph(f"Транскрибация: {len(session.get('transcript', ''))} символов")
    doc.add_paragraph(f"Моделей протестировано: {len(stats)}")

    if stats:
        doc.add_heading("Технические метрики", level=1)
        headers = ["Модель", "Успешность", "Ср. время, с", "Мин/Медиана/Макс, с",
                   "Ток/сек", "Токены вх→вых", "Прогоны"]
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Light Grid Accent 1"
        for idx, h in enumerate(headers):
            table.rows[0].cells[idx].paragraphs[0].add_run(h).bold = True
        for s in stats:
            row = table.add_row().cells
            row[0].text = f"{s['provider']}/{s['model']}"
            row[1].text = f"{s['success_rate']}%"
            row[2].text = str(s["avg_latency_sec"])
            row[3].text = f"{s['min_latency_sec']}/{s['median_latency_sec']}/{s['max_latency_sec']}"
            row[4].text = str(s["avg_tokens_per_sec"])
            row[5].text = f"{s['avg_tokens_in']}→{s['avg_tokens_out']}"
            row[6].text = f"{s['runs_ok']}/{s['runs_total']}"
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(10)
        doc.add_paragraph()

    report = session.get("report", "")
    if report.strip():
        doc.add_heading("Отчёт судьи", level=1)
        _add_markdown(doc, report)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
