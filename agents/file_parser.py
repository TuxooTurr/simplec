"""
Парсер файлов с валидацией безопасности.
Поддержка: PDF, Word, Excel, XML, изображения (OCR), TXT.
"""

import io

MAX_FILE_SIZE = 50 * 1024 * 1024  # 50 MB
ALLOWED_EXTENSIONS = {
    ".pdf", ".docx", ".doc", ".xlsx", ".xls",
    ".xml", ".png", ".jpg", ".jpeg", ".txt"
}


def validate_file(data: bytes, filename: str) -> None:
    """Валидация файла перед парсингом."""
    if len(data) > MAX_FILE_SIZE:
        raise ValueError(
            "Файл слишком большой: "
            + str(len(data) // (1024 * 1024))
            + " MB (макс. 50 MB)"
        )

    ext = ""
    if "." in filename:
        ext = "." + filename.rsplit(".", 1)[-1].lower()

    if ext not in ALLOWED_EXTENSIONS:
        raise ValueError(
            "Недопустимый формат: " + ext
            + " (разрешены: " + ", ".join(sorted(ALLOWED_EXTENSIONS)) + ")"
        )

    if len(data) == 0:
        raise ValueError("Файл пустой")


def parse_file(data: bytes, filename: str) -> str:
    """Извлекает текст из файла с предварительной валидацией."""

    validate_file(data, filename)

    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if ext == "pdf":
        return _parse_pdf(data)
    elif ext in ("docx", "doc"):
        return _parse_docx(data)
    elif ext in ("xlsx", "xls"):
        return _parse_xlsx(data)
    elif ext == "xml":
        return _parse_xml(data)
    elif ext in ("png", "jpg", "jpeg"):
        return _parse_image(data)
    elif ext == "txt":
        return _parse_txt(data)
    else:
        return data.decode("utf-8", errors="ignore")


def _parse_pdf(data: bytes) -> str:
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(io.BytesIO(data))
        pages = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                pages.append(
                    "--- Страница " + str(i + 1) + " ---\n" + text
                )
        result = "\n\n".join(pages)
        if not result.strip():
            return "[PDF: текст не извлечён, возможно скан]"
        return result
    except ImportError:
        return "[Ошибка: установите PyPDF2]"
    except Exception as e:
        return "[Ошибка PDF: " + str(e) + "]"


def _parse_docx(data: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(data))
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                cells = []
                for cell in row.cells:
                    if cell.text.strip():
                        cells.append(cell.text.strip())
                if cells:
                    paragraphs.append(" | ".join(cells))

        result = "\n".join(paragraphs)
        if not result.strip():
            return "[DOCX: документ пустой]"
        return result
    except ImportError:
        return "[Ошибка: установите python-docx]"
    except Exception as e:
        return "[Ошибка DOCX: " + str(e) + "]"


def _parse_xlsx(data: bytes) -> str:
    try:
        from openpyxl import load_workbook
        wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        sheets = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            for row in ws.iter_rows(values_only=True):
                cells = []
                for cell in row:
                    if cell is not None:
                        cells.append(str(cell))
                if cells:
                    rows.append(" | ".join(cells))
            if rows:
                sheets.append(
                    "=== Лист: " + sheet_name + " ===\n"
                    + "\n".join(rows)
                )
        wb.close()
        result = "\n\n".join(sheets)
        if not result.strip():
            return "[XLSX: таблица пустая]"
        return result
    except ImportError:
        return "[Ошибка: установите openpyxl]"
    except Exception as e:
        return "[Ошибка XLSX: " + str(e) + "]"


def _parse_xml(data: bytes) -> str:
    try:
        text = data.decode("utf-8", errors="ignore")
        if not text.strip():
            return "[XML: файл пустой]"
        return text
    except Exception as e:
        return "[Ошибка XML: " + str(e) + "]"


def _parse_image(data: bytes) -> str:
    try:
        from PIL import Image
        import pytesseract
        img = Image.open(io.BytesIO(data))
        text = pytesseract.image_to_string(img, lang="rus+eng")
        if not text.strip():
            return "[Изображение: текст не распознан]"
        return text
    except ImportError:
        return "[OCR недоступен: установите Pillow и pytesseract]"
    except Exception as e:
        return "[Ошибка OCR: " + str(e) + "]"


def _parse_txt(data: bytes) -> str:
    try:
        for encoding in ("utf-8", "cp1251", "latin-1"):
            try:
                text = data.decode(encoding)
                if text.strip():
                    return text
            except UnicodeDecodeError:
                continue
        return "[TXT: не удалось декодировать]"
    except Exception as e:
        return "[Ошибка TXT: " + str(e) + "]"

